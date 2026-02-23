"""Integrity verification for fail-closed restoration."""

from __future__ import annotations

import csv
import hashlib
from io import StringIO
from pathlib import Path
from typing import Iterable

from cowork_shield.models import EntityMapping
from cowork_shield.tokenizer.patterns import ANY_TOKEN_PATTERN

TOKEN_PATTERN = ANY_TOKEN_PATTERN


def compute_sha256(file_path: Path) -> str:
    """Compute SHA-256 hex digest of a file."""
    h = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


class IntegrityVerifier:
    """Verification checks for fail-closed restoration.

    Three levels:
    1. HMAC verification: every mapping's integrity tag is valid
    2. Completeness check: no tokens remain in restored output
    3. Hash comparison: restored file matches original (when available)
    """

    def __init__(self, token_generator):
        self._generator = token_generator

    def verify_all_hmacs(
        self,
        mappings: dict[str, EntityMapping],
    ) -> list[str]:
        """Verify HMAC tags on all mappings. Returns list of failed token texts."""
        failures = []
        for mapping in mappings.values():
            if not self._generator.verify_token(
                mapping.token, mapping.original_value
            ):
                failures.append(mapping.token.token_text)
        return failures

    def scan_for_remaining_tokens(
        self,
        file_path: Path,
        known_tokens: Iterable[str],
    ) -> list[str]:
        """Scan a restored file for any remaining unreplaced tokens."""
        all_text = self._extract_all_text(file_path)
        known_set = set(known_tokens)
        return self.scan_remaining_tokens_in_text(all_text, known_set)

    def scan_remaining_tokens_in_text(
        self,
        text: str,
        known_tokens: Iterable[str],
    ) -> list[str]:
        """Find unresolved token-pattern strings using one pass over text."""
        known_set = set(known_tokens)
        found = {match.group(0) for match in TOKEN_PATTERN.finditer(text)}
        if not found:
            return []

        remaining = []
        for token_text in sorted(found):
            if self._is_known_token_match(token_text, known_set):
                remaining.append(token_text)
                continue
            remaining.append(token_text)
        return remaining

    def extract_all_text(self, file_path: Path) -> str:
        """Extract text content for external token analysis."""
        return self._extract_all_text(file_path)

    @staticmethod
    def extract_token_matches(text: str) -> set[str]:
        """Extract all token-shaped strings from text in a single regex pass."""
        if not text:
            return set()
        return {match.group(0) for match in TOKEN_PATTERN.finditer(text)}

    @staticmethod
    def resolve_known_tokens(
        observed_tokens: Iterable[str],
        known_tokens: Iterable[str],
    ) -> set[str]:
        """Map observed tokens to canonical known token keys (supports legacy/v2 mix)."""
        known = set(known_tokens)
        resolved: set[str] = set()
        for token in observed_tokens:
            if token in known:
                resolved.add(token)
                continue
            if token.startswith("[") and token.endswith("]"):
                inner = token[1:-1]
                if inner in known:
                    resolved.add(inner)
            else:
                wrapped = f"[{token}]"
                if wrapped in known:
                    resolved.add(wrapped)
        return resolved

    def verify_hmacs_for_token_subset(
        self,
        mappings: dict[str, EntityMapping],
        token_subset: set[str],
    ) -> list[str]:
        """Verify HMAC integrity only for mappings required by this restore input."""
        if not token_subset:
            return []
        failures = []
        for mapping in mappings.values():
            token_text = mapping.token.token_text
            if token_text not in token_subset:
                continue
            if not self._generator.verify_token(mapping.token, mapping.original_value):
                failures.append(token_text)
        return failures

    def _extract_all_text(self, file_path: Path) -> str:
        """Extract all text content from a file for token scanning."""
        suffix = file_path.suffix.lower()

        if suffix == ".csv":
            return self._extract_csv_text(file_path)
        elif suffix == ".xlsx":
            return self._extract_xlsx_text(file_path)
        elif suffix == ".docx":
            return self._extract_docx_text(file_path)
        else:
            # Fall back to reading as text
            return file_path.read_text(encoding="utf-8", errors="replace")

    @staticmethod
    def _extract_csv_text(file_path: Path) -> str:
        """Extract all cell values from a CSV file."""
        text = file_path.read_text(encoding="utf-8-sig")
        parts = []
        reader = csv.reader(StringIO(text))
        for row in reader:
            parts.extend(row)
        return " ".join(parts)

    @staticmethod
    def _extract_xlsx_text(file_path: Path) -> str:
        """Extract all cell values from an xlsx file."""
        from openpyxl import load_workbook

        wb = load_workbook(str(file_path), data_only=False, read_only=True)
        parts = []
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        parts.append(str(cell.value))
        wb.close()
        return " ".join(parts)

    @staticmethod
    def _extract_docx_text(file_path: Path) -> str:
        """Extract all text from a docx file."""
        from docx import Document

        doc = Document(str(file_path))
        parts = []
        for para in doc.paragraphs:
            parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return " ".join(parts)

    @staticmethod
    def _token_present(all_text: str, token: str) -> bool:
        if token in all_text:
            return True

        if token.startswith("[") and token.endswith("]"):
            return token[1:-1] in all_text

        return f"[{token}]" in all_text

    @staticmethod
    def _is_known_token_match(token_text: str, known_tokens: set[str]) -> bool:
        if token_text in known_tokens:
            return True
        if token_text.startswith("[") and token_text.endswith("]"):
            return token_text[1:-1] in known_tokens
        return f"[{token_text}]" in known_tokens
