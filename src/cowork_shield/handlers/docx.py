"""Word .docx file handler with paragraph-level detection and run-level formatting."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

from cowork_shield.detection.engine import DetectionEngine
from cowork_shield.models import FileRecord, ReplacementRecord, now_iso
from cowork_shield.tokenizer.generator import TokenGenerator
from cowork_shield.tokenizer.replacer import TextReplacer
from cowork_shield.verification.verifier import compute_sha256


class DocxHandler:
    """Handles .docx files using python-docx.

    THE CRITICAL PROBLEM: Word splits text across runs unpredictably.
    "John Smith" might be stored as runs ["Jo", "hn Smi", "th"].

    SOLUTION: Paragraph-level detection with run-level redistribution.
    1. Concatenate all run texts → full paragraph string
    2. Detect PII in the full string
    3. Apply token replacements
    4. Redistribute modified text across runs, preserving formatting
    """

    def __init__(self):
        self._replacer = TextReplacer()

    @staticmethod
    def can_handle(file_path: Path) -> bool:
        return file_path.suffix.lower() == ".docx"

    def anonymize(
        self,
        input_path: Path,
        output_path: Path,
        detection_engine: DetectionEngine,
        token_generator: TokenGenerator,
        source_file: str = "",
    ) -> tuple[list[ReplacementRecord], FileRecord]:
        doc = Document(str(input_path))
        all_records: list[ReplacementRecord] = []
        total_entities = 0

        # Process body paragraphs
        for i, para in enumerate(doc.paragraphs):
            entities_found, records = self._process_paragraph(
                para, detection_engine, token_generator,
                f"para:{i}", source_file,
            )
            total_entities += entities_found
            all_records.extend(records)

        # Process tables
        for t, table in enumerate(doc.tables):
            for r, row in enumerate(table.rows):
                for c, cell in enumerate(row.cells):
                    for p, para in enumerate(cell.paragraphs):
                        entities_found, records = self._process_paragraph(
                            para, detection_engine, token_generator,
                            f"table:{t},row:{r},col:{c},para:{p}", source_file,
                        )
                        total_entities += entities_found
                        all_records.extend(records)

        # Process headers/footers
        for s, section in enumerate(doc.sections):
            for label, hf in [("header", section.header), ("footer", section.footer)]:
                if hf is None or hf.is_linked_to_previous:
                    continue
                for p, para in enumerate(getattr(hf, "paragraphs", [])):
                    entities_found, records = self._process_paragraph(
                        para, detection_engine, token_generator,
                        f"section:{s},{label},para:{p}", source_file,
                    )
                    total_entities += entities_found
                    all_records.extend(records)

        doc.save(str(output_path))

        file_record = FileRecord(
            file_path=str(input_path),
            file_hash_before=compute_sha256(input_path),
            file_hash_after=compute_sha256(output_path),
            anonymized_path=str(output_path),
            entities_found=total_entities,
            tokens_applied=len(all_records),
            timestamp=now_iso(),
            format="docx",
        )

        return all_records, file_record

    def restore(
        self,
        input_path: Path,
        output_path: Path,
        reverse_lookup: dict[str, str],
    ) -> None:
        doc = Document(str(input_path))

        # Process body paragraphs
        for para in doc.paragraphs:
            self._restore_paragraph(para, reverse_lookup)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._restore_paragraph(para, reverse_lookup)

        # Process headers/footers
        for section in doc.sections:
            for hf in [section.header, section.footer]:
                if hf is None:
                    continue
                for para in getattr(hf, "paragraphs", []):
                    self._restore_paragraph(para, reverse_lookup)

        doc.save(str(output_path))

    def _process_paragraph(
        self,
        paragraph: Paragraph,
        detection_engine: DetectionEngine,
        token_generator: TokenGenerator,
        para_id: str,
        source_file: str,
    ) -> tuple[int, list[ReplacementRecord]]:
        """Process a single paragraph: detect, replace, redistribute to runs."""
        runs = paragraph.runs
        if not runs:
            return 0, []

        full_text = "".join(run.text for run in runs)
        if not full_text.strip():
            return 0, []

        # Detect entities in full paragraph text
        entities = detection_engine.detect_in_cell(full_text, para_id)
        if not entities:
            return 0, []

        # Apply token replacements to get modified text
        new_text, records = self._replacer.replace_entities(
            full_text, entities, token_generator, source_file
        )

        # Redistribute modified text across runs, preserving formatting
        self._redistribute_to_runs(runs, full_text, new_text, entities)

        return len(entities), records

    def _redistribute_to_runs(
        self,
        runs: list,
        old_text: str,
        new_text: str,
        entities: list,
    ) -> None:
        """Redistribute modified text back into runs preserving formatting.

        Strategy:
        - Build a character-to-run-index mapping for the old text
        - Map each character in the new text to the formatting of the
          corresponding original character
        - Where a token replaces a span, all token chars get the formatting
          of the first run in the replaced span
        - Rebuild runs with correct text segments
        """
        if not runs:
            return

        # Build character-to-run-index mapping for old text
        char_to_run: list[int] = []
        for i, run in enumerate(runs):
            char_to_run.extend([i] * len(run.text))

        # Build offset mapping: for each position in new_text,
        # determine which run's formatting it should inherit
        #
        # We process replacements to build an old-position to new-position map.
        # Sort entities by start ascending for forward processing
        sorted_entities = sorted(entities, key=lambda e: e.start)

        # Build segments of new text with their run index assignments
        # Each segment is (text, run_index)
        new_char_to_run: list[int] = []
        old_pos = 0
        new_pos = 0

        for entity in sorted_entities:
            # Characters before this entity: unchanged, same run mapping
            before_len = entity.start - old_pos
            if before_len > 0:
                new_char_to_run.extend(char_to_run[old_pos : entity.start])
                new_pos += before_len

            # Find the token text that replaced this entity in new_text
            # The token length can be computed from the difference
            old_entity_len = entity.end - entity.start
            # We need to figure out how many chars the token occupies
            # in new_text at this position
            token = self._find_token_at_position(
                new_text, new_pos + before_len if new_char_to_run else new_pos,
                entity,
            )
            token_len = len(token) if token else old_entity_len

            # All token characters inherit formatting from the first char
            # of the replaced span
            first_run_idx = char_to_run[entity.start] if entity.start < len(char_to_run) else 0
            new_char_to_run.extend([first_run_idx] * token_len)

            old_pos = entity.end
            new_pos += before_len + token_len

        # Remaining characters after last entity
        if old_pos < len(old_text):
            new_char_to_run.extend(char_to_run[old_pos:])

        # Now rebuild runs: group consecutive characters by run index
        # to minimize the number of run modifications
        if not new_char_to_run:
            return

        # Assign text to each run
        run_texts: dict[int, list[str]] = {i: [] for i in range(len(runs))}

        # Build segments: consecutive chars with the same run_idx
        segments: list[tuple[int, str]] = []
        current_run_idx = new_char_to_run[0]
        current_chars = [new_text[0]] if new_text else []

        for i in range(1, len(new_text)):
            if i < len(new_char_to_run):
                run_idx = new_char_to_run[i]
            else:
                run_idx = current_run_idx

            if run_idx == current_run_idx:
                current_chars.append(new_text[i])
            else:
                segments.append((current_run_idx, "".join(current_chars)))
                current_run_idx = run_idx
                current_chars = [new_text[i]]

        if current_chars:
            segments.append((current_run_idx, "".join(current_chars)))

        # Assign segments to runs, clearing all runs first
        for run in runs:
            run.text = ""

        # Group segments by run index to handle multiple segments per run
        for run_idx, text in segments:
            if run_idx < len(runs):
                runs[run_idx].text += text
            elif runs:
                # Fallback: append to last run
                runs[-1].text += text

    def _find_token_at_position(self, text: str, pos: int, entity) -> str | None:
        """Find the token text at a given position in the new text."""
        # Token format is PREFIX_NNN, e.g., PERSON_001
        prefix = entity.entity_type.token_prefix
        # Look for the token starting at roughly the expected position
        # We search a window around the position since offsets may shift
        search_start = max(0, pos - 5)
        search_end = min(len(text), pos + len(prefix) + 10)
        window = text[search_start:search_end]

        import re
        pattern = rf"{re.escape(prefix)}_\d{{3,5}}"
        match = re.search(pattern, window)
        if match:
            return match.group(0)
        return None

    def _restore_paragraph(
        self,
        paragraph: Paragraph,
        reverse_lookup: dict[str, str],
    ) -> None:
        """Restore tokens in a paragraph back to original values."""
        runs = paragraph.runs
        if not runs:
            return

        full_text = "".join(run.text for run in runs)
        restored = self._replacer.restore_tokens(full_text, reverse_lookup)

        if restored == full_text:
            return

        # Simple redistribution for restoration: put all text in first run,
        # clear others. This is simpler than anonymization because we don't
        # need to preserve per-entity formatting.
        if runs:
            runs[0].text = restored
            for run in runs[1:]:
                run.text = ""
