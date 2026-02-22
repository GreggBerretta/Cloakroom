"""PDF input-only handler that extracts markdown before anonymization."""

from __future__ import annotations

from pathlib import Path

from cowork_shield.detection.engine import DetectionEngine
from cowork_shield.exceptions import CoWorkShieldError, PdfInputOnlyError
from cowork_shield.extractors.pdf_markdown import PDFExtractor
from cowork_shield.models import DetectedEntity, FileRecord, ReplacementRecord, now_iso
from cowork_shield.tokenizer.generator import TokenGenerator
from cowork_shield.tokenizer.replacer import TextReplacer
from cowork_shield.verification.verifier import compute_sha256


class PdfHandler:
    """Converts PDF input to markdown/docx and anonymizes extracted text."""

    _OUTPUT_FORMATS = {"md", "docx"}

    def __init__(self, pdf_output_format: str = "md", extractor: PDFExtractor | None = None):
        output_format = (pdf_output_format or "md").strip().lower()
        if output_format not in self._OUTPUT_FORMATS:
            raise CoWorkShieldError(
                f"Invalid PDF output format: {output_format}. Use one of: md, docx."
            )
        self._pdf_output_format = output_format
        self._extractor = extractor or PDFExtractor()
        self._replacer = TextReplacer()

    @staticmethod
    def can_handle(file_path: Path) -> bool:
        return file_path.suffix.lower() == ".pdf"

    def anonymize(
        self,
        input_path: Path,
        output_path: Path,
        detection_engine: DetectionEngine,
        token_generator: TokenGenerator,
        source_file: str = "",
        language: str = "auto",
    ) -> tuple[list[ReplacementRecord], FileRecord]:
        extraction = self._extractor.extract(input_path)
        entities = _detect_entities(
            detection_engine,
            text=extraction.markdown,
            source_id="pdf:0",
            language=language,
        )
        anonymized_text, records = self._replacer.replace_entities(
            extraction.markdown,
            entities,
            token_generator,
            source_file=source_file,
        )

        output_suffix = output_path.suffix.lower().lstrip(".")
        if output_suffix not in self._OUTPUT_FORMATS:
            raise CoWorkShieldError(
                "PDF anonymize output must be .md or .docx. "
                f"Received: {output_path.suffix or '<none>'}"
            )

        if output_suffix == "docx":
            _write_markdown_to_docx(anonymized_text, output_path)
        else:
            output_path.write_text(anonymized_text, encoding="utf-8")

        file_record = FileRecord(
            file_path=str(input_path),
            file_hash_before=compute_sha256(input_path),
            file_hash_after=compute_sha256(output_path),
            anonymized_path=str(output_path),
            entities_found=len(entities),
            tokens_applied=len(records),
            timestamp=now_iso(),
            format=f"pdf->{output_suffix}",
            applied_tokens=sorted({record.token_text for record in records}),
        )
        # Keep backend provenance in override events for troubleshooting/audit breadcrumbs.
        file_record.override_events = [f"pdf_extractor:{extraction.backend}"]
        return records, file_record

    def restore(
        self,
        input_path: Path,
        output_path: Path,
        reverse_lookup: dict[str, str],
    ) -> None:
        raise PdfInputOnlyError()


def _write_markdown_to_docx(markdown_text: str, output_path: Path) -> None:
    from docx import Document

    doc = Document()
    lines = markdown_text.splitlines()
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)

    doc.save(output_path)


def _detect_entities(
    detection_engine: DetectionEngine,
    *,
    text: str,
    source_id: str,
    language: str,
) -> list[DetectedEntity]:
    try:
        return detection_engine.detect_in_cell(text, source_id, language=language)
    except TypeError:
        # Compatibility for tests using stub engines without language arg.
        return detection_engine.detect_in_cell(text, source_id)
