"""Shared UI-facing API wrappers for TUI and Gradio frontends."""

from __future__ import annotations

from dataclasses import dataclass
import html
from pathlib import Path

from cowork_shield.detection.engine import DetectionEngine
from cowork_shield.exceptions import (
    DetectionError,
    HallucinationDetectedError,
    IncompleteRestorationError,
    IntegrityError,
    KeychainError,
    ModelHashMismatchError,
    PdfExtractionError,
    PdfInputOnlyError,
    ReplayMismatchError,
    UnsupportedFormatError,
    WorkspaceExpiredError,
    WorkspaceNotFoundError,
    XLSXContentLossRiskError,
)
from cowork_shield.pipeline.anonymize import AnonymizePipeline
from cowork_shield.pipeline.restore import RestorePipeline
from cowork_shield.workspace.manager import WorkspaceManager


@dataclass(frozen=True)
class UIOperationResult:
    """Result payload returned to UI frontends."""

    path: str
    summary: str
    entity_rows: list[dict[str, str]]
    entity_table_html: str


def sanitize_ui_error(exc: Exception) -> tuple[str, str]:
    """Return a stable, non-sensitive UI-safe error code/message pair."""
    code = exc.__class__.__name__

    if isinstance(exc, UnsupportedFormatError):
        return code, "Unsupported format. Use PDF, CSV, XLSX, DOCX, TXT, or MD."
    if isinstance(exc, PdfExtractionError):
        return code, "PDF extraction failed. Install Docling or PyMuPDF and retry."
    if isinstance(exc, PdfInputOnlyError):
        return (
            code,
            "PDF is input-only. Restore from tokenized Markdown (.md) or DOCX (.docx).",
        )
    if isinstance(exc, DetectionError):
        return code, "PII detection failed. Verify language model installation and retry."
    if isinstance(exc, WorkspaceNotFoundError):
        return code, "Workspace not found. Select an existing workspace."
    if isinstance(exc, WorkspaceExpiredError):
        return code, "Workspace expired. Create a new workspace and retry."
    if isinstance(exc, XLSXContentLossRiskError):
        return (
            code,
            "XLSX includes charts/images that may be dropped. Enable lossy XLSX only with explicit confirmation.",
        )
    if isinstance(exc, ModelHashMismatchError):
        return code, "Model hash mismatch. Deterministic safeguards blocked this operation."
    if isinstance(exc, ReplayMismatchError):
        return code, "Deterministic replay mismatch detected. Operation aborted."
    if isinstance(exc, HallucinationDetectedError):
        return code, "Potential AI mutation detected. Restore aborted."
    if isinstance(exc, IntegrityError):
        return code, "Vault integrity verification failed. Restore aborted."
    if isinstance(exc, IncompleteRestorationError):
        return code, "Restore incomplete. Unresolved tokens remain in output."
    if isinstance(exc, KeychainError):
        return code, "Keychain operation failed. Verify Keychain access and retry."
    if isinstance(exc, OSError):
        return code, "Filesystem operation failed. Verify path permissions and free disk space."
    return code, "Operation failed. See logs for details."


def get_workspaces() -> list[str]:
    """Return known workspace names for UI selectors."""
    mgr = WorkspaceManager()
    names = [ws["name"] for ws in mgr.list_workspaces()]
    if "default" not in names:
        names.insert(0, "default")
    return sorted(set(names), key=lambda n: (n != "default", n.lower()))


def preview_entities(
    file_path: str | Path,
    *,
    score_threshold: float = 0.7,
    language: str = "auto",
    max_rows: int = 200,
) -> list[dict[str, str]]:
    """Detect and return entity rows for attestation/review."""
    path = Path(file_path).expanduser().resolve()
    text = _read_supported_text(path)
    detection = DetectionEngine(score_threshold=score_threshold)
    entities = detection.detect(text, language=language)

    rows: list[dict[str, str]] = []
    for entity in entities[:max_rows]:
        rows.append(
            {
                "type": entity.entity_type.value,
                "text": entity.text,
                "start": str(entity.start),
                "end": str(entity.end),
                "score": f"{entity.score:.3f}",
            }
        )
    return rows


def anonymize_file(
    file_path: str | Path,
    workspace: str,
    *,
    output_path: str | Path | None = None,
    ttl_hours: int = 168,
    score_threshold: float = 0.7,
    language: str = "auto",
    allow_lossy_xlsx: bool = False,
    pdf_output_format: str = "md",
    force_reanonymize: bool = False,
    reason: str = "",
) -> UIOperationResult:
    """Run anonymization and return UI-friendly payload."""
    mgr = WorkspaceManager()
    ctx = mgr.get_or_create_workspace(workspace, ttl_hours=ttl_hours)

    input_path = Path(file_path).expanduser().resolve()
    out_path = Path(output_path).expanduser().resolve() if output_path else None
    entity_rows = preview_entities(
        input_path,
        score_threshold=score_threshold,
        language=language,
    )

    pipeline = AnonymizePipeline(
        ctx,
        score_threshold=score_threshold,
        language=language,
        force_reanonymize=force_reanonymize,
        override_reason=reason,
        override_user="ui",
        allow_lossy_xlsx=allow_lossy_xlsx,
        pdf_output_format=pdf_output_format,
    )
    result = pipeline.run(input_path, out_path)

    summary = (
        f"Anonymized {result.input_path.name} -> {result.output_path.name}. "
        f"Entities: {result.entities_found}. Tokens: {result.tokens_applied}."
    )
    if input_path.suffix.lower() == ".pdf":
        summary += (
            " PDF is input-only: output is extracted Markdown/DOCX; "
            "the original PDF binary is not reconstructed."
        )
    return UIOperationResult(
        path=str(result.output_path),
        summary=summary,
        entity_rows=entity_rows,
        entity_table_html=render_entity_table(entity_rows),
    )


def restore_file(
    file_path: str | Path,
    workspace: str,
    *,
    output_path: str | Path | None = None,
) -> UIOperationResult:
    """Run restoration and return UI-friendly payload."""
    mgr = WorkspaceManager()
    ctx = mgr.get_active_workspace(workspace)

    input_path = Path(file_path).expanduser().resolve()
    out_path = Path(output_path).expanduser().resolve() if output_path else None
    pipeline = RestorePipeline(ctx)
    result = pipeline.run(input_path, out_path)

    summary = (
        f"Restored {result.input_path.name} -> {result.output_path.name}. "
        f"Tokens restored: {result.tokens_restored}."
    )
    return UIOperationResult(
        path=str(result.output_path),
        summary=summary,
        entity_rows=[],
        entity_table_html=render_entity_table([]),
    )


def render_entity_table(rows: list[dict[str, str]]) -> str:
    """Render an HTML table for entity attestation in web UI."""
    if not rows:
        return "<p><strong>No entities detected.</strong></p>"

    header = (
        "<thead><tr>"
        "<th>Type</th><th>Text</th><th>Start</th><th>End</th><th>Score</th>"
        "</tr></thead>"
    )
    body_rows = []
    for row in rows:
        body_rows.append(
            "<tr>"
            f"<td>{html.escape(row['type'])}</td>"
            f"<td>{html.escape(row['text'])}</td>"
            f"<td>{html.escape(row['start'])}</td>"
            f"<td>{html.escape(row['end'])}</td>"
            f"<td>{html.escape(row['score'])}</td>"
            "</tr>"
        )
    body = "<tbody>" + "".join(body_rows) + "</tbody>"
    return (
        "<table border='1' cellpadding='6' cellspacing='0' style='border-collapse:collapse;'>"
        + header
        + body
        + "</table>"
    )


def _read_supported_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="replace")
    if suffix == ".md":
        return path.read_text(encoding="utf-8", errors="replace")
    if suffix == ".csv":
        return path.read_text(encoding="utf-8-sig", errors="replace")
    if suffix == ".pdf":
        from cowork_shield.extractors.pdf_markdown import extract_pdf_to_markdown

        return extract_pdf_to_markdown(path)
    if suffix == ".docx":
        from docx import Document

        doc = Document(str(path))
        parts = [para.text for para in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    if suffix == ".xlsx":
        from openpyxl import load_workbook

        wb = load_workbook(str(path), data_only=False, read_only=True)
        parts: list[str] = []
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        parts.append(str(cell.value))
        wb.close()
        return "\n".join(parts)

    raise UnsupportedFormatError(suffix)
