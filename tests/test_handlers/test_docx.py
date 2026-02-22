"""Tests for the Word .docx file handler."""

import os
from pathlib import Path

import pytest
from docx import Document

from cowork_shield.detection.engine import DetectionEngine
from cowork_shield.handlers.docx import DocxHandler
from cowork_shield.tokenizer.generator import TokenGenerator

FIXTURES_DIR = Path(__file__).parent.parent / "fixtures"


@pytest.fixture(scope="module")
def detection_engine():
    return DetectionEngine(score_threshold=0.5)


@pytest.fixture
def token_generator():
    return TokenGenerator(os.urandom(32))


class TestDocxHandler:
    def test_anonymize_basic(self, tmp_path, detection_engine, token_generator):
        handler = DocxHandler()
        input_path = FIXTURES_DIR / "sample_report.docx"
        output_path = tmp_path / "anonymized.docx"

        records, file_record = handler.anonymize(
            input_path, output_path, detection_engine, token_generator
        )

        assert output_path.exists()
        assert file_record.format == "docx"
        assert file_record.entities_found > 0
        assert len(records) > 0

        # Verify tokens are in the output
        doc = Document(str(output_path))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "PERSON_" in all_text or "ORG_" in all_text or "EMAIL_" in all_text

    def test_anonymize_tables(self, tmp_path, detection_engine, token_generator):
        handler = DocxHandler()
        input_path = FIXTURES_DIR / "sample_with_tables.docx"
        output_path = tmp_path / "anonymized.docx"

        records, file_record = handler.anonymize(
            input_path, output_path, detection_engine, token_generator
        )

        assert file_record.entities_found > 0

        # Verify table cells are anonymized
        doc = Document(str(output_path))
        table_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + " "

        # Should contain tokens
        has_tokens = any(
            prefix in table_text
            for prefix in ["PERSON_", "EMAIL_", "PHONE_"]
        )
        assert has_tokens

    def test_round_trip_paragraphs(self, tmp_path, detection_engine, token_generator):
        handler = DocxHandler()
        input_path = FIXTURES_DIR / "sample_report.docx"
        anon_path = tmp_path / "anonymized.docx"
        restored_path = tmp_path / "restored.docx"

        handler.anonymize(input_path, anon_path, detection_engine, token_generator)

        reverse_lookup = token_generator.get_reverse_lookup()
        handler.restore(anon_path, restored_path, reverse_lookup)

        # Compare paragraph texts
        doc_orig = Document(str(input_path))
        doc_rest = Document(str(restored_path))

        orig_texts = [p.text for p in doc_orig.paragraphs if p.text.strip()]
        rest_texts = [p.text for p in doc_rest.paragraphs if p.text.strip()]

        assert len(orig_texts) == len(rest_texts)
        for orig, rest in zip(orig_texts, rest_texts):
            assert orig == rest, f"Paragraph mismatch:\n  orig: {orig}\n  rest: {rest}"

    def test_can_handle(self):
        assert DocxHandler.can_handle(Path("report.docx"))
        assert DocxHandler.can_handle(Path("report.DOCX"))
        assert not DocxHandler.can_handle(Path("report.xlsx"))

    def test_simple_paragraph_anonymize(self, tmp_path, detection_engine, token_generator):
        """Test with a simple document we create ourselves."""
        handler = DocxHandler()

        # Create a simple doc
        doc = Document()
        doc.add_paragraph("John Smith works at Acme Corporation.")
        input_path = tmp_path / "simple.docx"
        doc.save(str(input_path))

        output_path = tmp_path / "simple_anon.docx"
        records, _ = handler.anonymize(
            input_path, output_path, detection_engine, token_generator
        )

        # Verify anonymization
        doc_anon = Document(str(output_path))
        text = doc_anon.paragraphs[0].text
        assert "John Smith" not in text
        assert "PERSON_" in text

    def test_empty_document(self, tmp_path, detection_engine, token_generator):
        handler = DocxHandler()

        doc = Document()
        doc.add_paragraph("")
        input_path = tmp_path / "empty.docx"
        doc.save(str(input_path))

        output_path = tmp_path / "empty_anon.docx"
        records, file_record = handler.anonymize(
            input_path, output_path, detection_engine, token_generator
        )

        assert file_record.entities_found == 0
        assert len(records) == 0
